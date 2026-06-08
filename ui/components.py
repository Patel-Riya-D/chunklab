from __future__ import annotations

import base64
from io import BytesIO
import html as html_lib
import re

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components

from utils.models import Chunk, DocumentType, ExtractedUnit, RetrievalResult, UploadedDocument
from utils.text import estimate_tokens


_CODE_LANGUAGES = {
    ".py": "python",
    ".js": "javascript",
    ".ts": "typescript",
    ".tsx": "tsx",
    ".jsx": "jsx",
    ".java": "java",
    ".go": "go",
    ".rs": "rust",
    ".cpp": "cpp",
    ".c": "c",
    ".h": "c",
    ".cs": "csharp",
    ".php": "php",
    ".rb": "ruby",
}


def _decode_document_text(document: UploadedDocument) -> str:
    return document.content.decode("utf-8", errors="ignore")


def _docx_text(content: bytes) -> str:
    try:
        from docx import Document
    except Exception:
        return ""

    try:
        doc = Document(BytesIO(content))
    except Exception:
        return ""

    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    table_lines = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_lines.append(" | ".join(cells))
    return "\n\n".join([*paragraphs, *table_lines]).strip()


def show_original_document(document: UploadedDocument) -> None:
    with st.expander("Original Document Preview", expanded=True):
        st.download_button(
            "Download original",
            data=document.content,
            file_name=document.name,
            mime="application/octet-stream",
            use_container_width=True,
        )

        if document.doc_type == DocumentType.PDF:
            encoded_pdf = base64.b64encode(document.content).decode("ascii")
            st.markdown(
                f"""
                <iframe
                    src="data:application/pdf;base64,{encoded_pdf}"
                    width="100%"
                    height="760"
                    style="border: 1px solid rgba(128, 128, 128, 0.35); border-radius: 8px;"
                ></iframe>
                """,
                unsafe_allow_html=True,
            )
            return

        if document.doc_type == DocumentType.DOCX:
            preview_text = _docx_text(document.content)
            st.caption("DOCX is shown as readable text; use download to open the original formatted file.")
            if preview_text:
                st.text_area("DOCX text preview", preview_text, height=520, label_visibility="collapsed")
            else:
                st.info("No readable DOCX text preview could be generated.")
            return

        text = _decode_document_text(document)
        if document.doc_type == DocumentType.MD:
            st.markdown(text)
        elif document.doc_type == DocumentType.HTML:
            components.html(text, height=760, scrolling=True)
        elif document.doc_type == DocumentType.CODE:
            st.code(text, language=_CODE_LANGUAGES.get(document.suffix, "text"))
        else:
            escaped = html_lib.escape(text)
            st.markdown(
                f"""
                <div style="max-height: 760px; overflow: auto; white-space: pre-wrap;
                            border: 1px solid rgba(128, 128, 128, 0.35);
                            border-radius: 8px; padding: 1rem;">
                    {escaped}
                </div>
                """,
                unsafe_allow_html=True,
            )


def _table_dataframe(rows: list[list[str]]) -> pd.DataFrame:
    if not rows:
        return pd.DataFrame()
    header, body = rows[0], rows[1:]
    column_names = [name or f"Column {index}" for index, name in enumerate(header, start=1)]
    return pd.DataFrame(body, columns=column_names)


def _show_table_rows(rows: object) -> bool:
    if not isinstance(rows, list) or not rows:
        return False
    if not all(isinstance(row, list) for row in rows):
        return False
    st.table(_table_dataframe(rows))
    return True


def _table_rows_from_text(text: str) -> list[list[str]]:
    rows = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped.startswith("|") or "|" not in stripped.strip("|"):
            continue
        cells = [cell.strip().replace("\\|", "|") for cell in stripped.strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells if cell):
            continue
        if len([cell for cell in cells if cell]) >= 2:
            rows.append(cells)
    if len(rows) < 2:
        return []
    column_count = max(len(row) for row in rows)
    return [row + [""] * (column_count - len(row)) for row in rows]


def _show_table_text(text: str) -> bool:
    rows = _table_rows_from_text(text)
    if not rows:
        return False
    st.table(_table_dataframe(rows))
    return True


def _show_text(unit_type: str, text: str) -> None:
    if unit_type in {"code_symbol", "code_context"}:
        st.code(text, language="python")
    else:
        st.write(text)


def show_units(units: list[ExtractedUnit], depth: int = 0) -> None:
    for unit in units:
        label = f"{unit.id} | {unit.unit_type}"
        if unit.metadata.get("title"):
            label += f" | {unit.metadata['title']}"
        with st.container(border=True):
            st.markdown(f"**{label}**")
            if unit.metadata:
                visible_metadata = {key: value for key, value in unit.metadata.items() if key != "rows"}
                st.caption(", ".join(f"{key}: {value}" for key, value in visible_metadata.items()))
            if not _show_table_rows(unit.metadata.get("rows")) and not _show_table_text(unit.text):
                _show_text(unit.unit_type, unit.text)
            if unit.children:
                show_units(unit.children, depth + 1)


def chunk_table(chunks: list[Chunk]) -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "id": chunk.id,
                "role": chunk.metadata.get("chunk_role", chunk.metadata.get("role", "")),
                "parent": chunk.parent_id or "",
                "children": len(chunk.children),
                "chars": len(chunk.text),
                "tokens_est": estimate_tokens(chunk.text),
                "overlap": chunk.overlap_with_previous,
                "source": chunk.metadata.get("source_unit", chunk.metadata.get("title", "")),
            }
            for chunk in chunks
        ]
    )


def _chunk_role(chunk: Chunk) -> str:
    role = str(chunk.metadata.get("chunk_role", chunk.metadata.get("role", ""))).lower()
    if role:
        return role
    if chunk.parent_id:
        return "child"
    if chunk.children:
        return "parent"
    return "chunk"


def _chunk_title(chunk: Chunk) -> str:
    return str(chunk.metadata.get("title") or chunk.metadata.get("section_path") or chunk.metadata.get("source_unit") or chunk.id)


def _page_label(chunk: Chunk) -> str:
    page = chunk.metadata.get("page")
    return f" | page {page}" if page else ""


def _relationship_summary(chunk: Chunk) -> str:
    role = _chunk_role(chunk)
    if role == "parent":
        child_ids = ", ".join(chunk.children[:4])
        if len(chunk.children) > 4:
            child_ids += f", +{len(chunk.children) - 4} more"
        return f"Parent chunk | children: {len(chunk.children)}" + (f" | {child_ids}" if child_ids else "")
    if role == "child":
        return f"Child chunk | parent: {chunk.parent_id or chunk.metadata.get('parent_chunk_id', '')}"
    return "Standalone chunk"


def _compact_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _format_score_details(details: dict[str, float | str], hidden_keys: set[str]) -> str:
    return ", ".join(f"{key}: {value}" for key, value in details.items() if key not in hidden_keys)


def _write_text_or_table(chunk: Chunk, text: str | None = None) -> None:
    content = chunk.text if text is None else text
    if not _show_table_rows(chunk.metadata.get("rows")) and not _show_table_text(content):
        st.write(_compact_text(content))


def show_chunk(chunk: Chunk, index: int | None = None, total: int | None = None) -> None:
    with st.container(border=True):
        role = _chunk_role(chunk)
        role_label = "Parent" if role == "parent" else "Child" if role == "child" else "Chunk"
        prefix = f"Chunk {index} of {total} - " if index is not None and total is not None else ""
        label = f"{prefix}{role_label.upper()} | {chunk.id}"
        st.markdown(f"**{label}**")
        st.caption(_relationship_summary(chunk))
        cols = st.columns(5)
        cols[0].metric("Role", role_label)
        cols[1].metric("Characters", len(chunk.text))
        cols[2].metric("Estimated tokens", estimate_tokens(chunk.text))
        cols[3].metric("Overlap", chunk.overlap_with_previous)
        cols[4].metric("Children", len(chunk.children))
        meta_parts = []
        if chunk.parent_id:
            meta_parts.append(f"parent: {chunk.parent_id}")
        if chunk.metadata.get("parent_section_title"):
            meta_parts.append(f"parent section: {chunk.metadata['parent_section_title']}")
        if chunk.metadata.get("section_path"):
            meta_parts.append(f"path: {chunk.metadata['section_path']}")
        if chunk.metadata.get("title"):
            meta_parts.append(f"section: {chunk.metadata['title']}")
        if chunk.metadata.get("page"):
            meta_parts.append(f"page: {chunk.metadata['page']}")
        if chunk.metadata.get("source_unit"):
            meta_parts.append(f"source: {chunk.metadata['source_unit']}")
        if meta_parts:
            st.caption(" | ".join(meta_parts))
        if not _show_table_rows(chunk.metadata.get("rows")) and not _show_table_text(chunk.text):
            _show_text("code_symbol" if chunk.strategy_id == "code_function" else "", chunk.text)


def show_retrieved(results: list[RetrievalResult]) -> None:
    hidden_keys = {
        "retrieved_role",
        "matched_child",
        "expanded_parent",
        "retrieval_level",
        "expansion_level",
        "expanded_parent_source",
        "expanded_parent_children",
        "expanded_parent_text",
        "expanded_parent_rank",
        "expanded_parent_score",
        "expanded_parent_score_details",
        "matched_parent",
        "parent_children",
    }
    visible_results = [
        result
        for result in results
        if str(result.score_details.get("retrieved_role") or _chunk_role(result.chunk)).lower() != "parent"
    ]
    for display_rank, result in enumerate(visible_results, start=1):
        chunk = result.chunk
        role = str(result.score_details.get("retrieved_role") or _chunk_role(chunk)).lower()
        title = _chunk_title(chunk)
        page_label = _page_label(chunk)
        role_label = "Parent" if role == "parent" else "Child" if role == "child" else "Chunk"
        with st.container(border=True):
            st.markdown(f"**Rank {display_rank}**")

            if role == "child" and result.score_details.get("expanded_parent"):
                st.markdown("**Matched Child Chunk**")
                st.caption(f"Chunk ID: {chunk.id}")
                _write_text_or_table(chunk)
                st.caption(f"Child Similarity Score: {result.score:.3f}")
                trace = f"retrieval level: child | expansion level: parent"
                if result.score_details.get("expanded_parent_source"):
                    trace += f" | parent source: {result.score_details['expanded_parent_source']}"
                st.caption(trace)
                details = _format_score_details(result.score_details, hidden_keys)
                if details:
                    st.caption(details)

                st.divider()
                st.markdown("**Expanded Parent Chunk**")
                st.caption(f"Parent ID: {result.score_details['expanded_parent']}")
                parent_text = str(result.score_details.get("expanded_parent_text", "")).strip()
                if parent_text:
                    st.write(_compact_text(parent_text))
                parent_trace = []
                parent_trace.append(f"expanded from child: {result.score_details.get('matched_child', chunk.id)}")
                if result.score_details.get("expanded_parent_children"):
                    parent_trace.append(f"children: {result.score_details['expanded_parent_children']}")
                st.caption(" | ".join(parent_trace))
            else:
                st.markdown(f"**{role_label} match | {title}{page_label}**")
                st.caption(f"Chunk ID: {chunk.id}")
                _write_text_or_table(chunk)
                st.caption(f"Score: {result.score:.3f}")
                trace_parts = []
                if result.score_details.get("matched_parent"):
                    trace_parts.append(f"matched parent: {result.score_details['matched_parent']}")
                if result.score_details.get("parent_children"):
                    trace_parts.append(f"children: {result.score_details['parent_children']}")
                if trace_parts:
                    st.caption(" | ".join(trace_parts))
                details = _format_score_details(result.score_details, hidden_keys)
                if details:
                    st.caption(details)
