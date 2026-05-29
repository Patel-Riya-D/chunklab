from __future__ import annotations

from abc import ABC, abstractmethod
from io import BytesIO
from typing import Iterable

from utils.models import ExtractedUnit, ExtractionResult, UploadedDocument
from utils.text import html_to_text, normalize_whitespace


class BaseExtractor(ABC):
    method_id: str
    name: str
    description: str

    @abstractmethod
    def extract(self, document: UploadedDocument) -> ExtractionResult:
        raise NotImplementedError

    def _read_text(self, document: UploadedDocument) -> str:
        if document.suffix == ".pdf":
            return "\n\n".join(self._read_pdf_pages(document.content))
        if document.suffix == ".docx":
            return self._read_docx(document.content)
        raw = document.content.decode("utf-8", errors="ignore")
        if document.suffix in {".html", ".htm"}:
            return html_to_text(raw)
        return raw

    def _read_pdf_pages(self, content: bytes) -> list[str]:
        try:
            import fitz

            doc = fitz.open(stream=content, filetype="pdf")
            return [page.get_text("text").strip() for page in doc]
        except Exception:
            try:
                import pdfplumber

                with pdfplumber.open(BytesIO(content)) as pdf:
                    return [(page.extract_text() or "").strip() for page in pdf.pages]
            except Exception:
                return ["PDF text extraction requires PyMuPDF or pdfplumber."]

    def _read_pdf_blocks(self, content: bytes) -> list[dict[str, object]]:
        try:
            import fitz

            doc = fitz.open(stream=content, filetype="pdf")
            blocks: list[dict[str, object]] = []
            for page_index, page in enumerate(doc, start=1):
                for block_index, block in enumerate(page.get_text("blocks"), start=1):
                    text = normalize_whitespace(str(block[4]))
                    if text:
                        blocks.append(
                            {
                                "text": text,
                                "page": page_index,
                                "block": block_index,
                                "bbox": tuple(round(float(v), 2) for v in block[:4]),
                            }
                        )
            return blocks
        except Exception:
            return []

    def _read_docx(self, content: bytes) -> str:
        try:
            from docx import Document

            doc = Document(BytesIO(content))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            tables = []
            for table in doc.tables:
                rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
                tables.append("\n".join(rows))
            return "\n\n".join(paragraphs + tables)
        except Exception:
            return content.decode("utf-8", errors="ignore")

    def _result(
        self,
        text: str,
        units: Iterable[ExtractedUnit],
        metadata: dict[str, object] | None = None,
    ) -> ExtractionResult:
        unit_list = list(units)
        return ExtractionResult(
            method_id=self.method_id,
            method_name=self.name,
            text=text,
            units=unit_list,
            metadata={
                "unit_count": len(unit_list),
                "character_count": len(text),
                **(metadata or {}),
            },
        )

