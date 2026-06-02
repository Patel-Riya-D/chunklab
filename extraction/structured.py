from __future__ import annotations

import re
import logging
from collections import Counter
from dataclasses import dataclass
from io import BytesIO

from extraction.base import BaseExtractor
from utils.models import ExtractedUnit, ExtractionResult, UploadedDocument
from utils.text import html_to_text


@dataclass
class _PdfTextBlock:
    text: str
    page: int
    bbox: tuple[float, float, float, float]
    max_size: float
    is_table: bool = False


def _section_units(text: str) -> list[ExtractedUnit]:
    top_level_sections = _top_level_section_units(text)
    if top_level_sections:
        return top_level_sections

    lines = text.splitlines()
    sections: list[ExtractedUnit] = []
    current_title = "Introduction"
    current_lines: list[str] = []
    current_level = 1

    def flush() -> None:
        if current_lines:
            idx = len(sections) + 1
            sections.append(
                ExtractedUnit(
                    id=f"section-{idx}",
                    text="\n".join(current_lines).strip(),
                    unit_type="section",
                    metadata={"title": current_title, "level": current_level, "section": idx},
                )
            )

    for line_index, line in enumerate(lines):
        next_line = next((item.strip() for item in lines[line_index + 1 :] if item.strip()), "")
        md = re.match(r"^(#{1,6})\s+(.+)$", line)
        numbered = _numbered_heading_match(line, next_line)
        inferred = _looks_like_text_heading(line, next_line)
        if inferred and not md and not numbered and not any(item.strip() for item in current_lines) and sections:
            current_lines = [line]
            continue
        if md or numbered or inferred:
            flush()
            current_lines = []
            if md:
                current_level = len(md.group(1))
                current_title = md.group(2).strip()
            else:
                current_level = numbered.group(1).count(".") + 1 if numbered else 1
                current_title = numbered.group(2).strip() if numbered else line.strip()
        else:
            current_lines.append(line)
    flush()
    return sections or [ExtractedUnit(id="section-1", text=text, unit_type="section", metadata={"title": "Document", "level": 1})]


def _numbered_heading_match(line: str, next_line: str = "") -> re.Match[str] | None:
    stripped = line.strip()
    match = re.match(r"^(\d+(?:\.\d+)*)[.)]?\s+(.{3,})$", stripped)
    if not match:
        return None
    title = match.group(2).strip()
    if title.startswith("**") or " - " in title or " – " in title or " — " in title:
        return None
    if len(title) > 90 or len(title.split()) > 8:
        return None
    if re.search(r"[.!?;:]$", title):
        return None
    if not next_line or len(next_line) < 25:
        return None
    return match


def _looks_like_text_heading(line: str, next_line: str = "") -> bool:
    stripped = line.strip()
    if not stripped or not re.search(r"[A-Za-z]", stripped):
        return False
    if _is_list_or_table_line(stripped):
        return False
    if len(stripped) > 90 or len(stripped.split()) > 8:
        return False
    if re.search(r"[.!?;:]$", stripped):
        return False
    if not next_line or len(next_line) < 25:
        return False
    words = re.findall(r"[A-Za-z][A-Za-z/&-]*", stripped)
    if not words:
        return False
    title_words = sum(1 for word in words if word[:1].isupper() or word.isupper())
    return title_words / len(words) >= 0.65


def _is_list_or_table_line(stripped: str) -> bool:
    return bool(
        re.match(r"^[-*+]\s+", stripped)
        or re.match(r"^\d+[.)]\s+", stripped)
        or re.match(r"^[-*_]{3,}$", stripped)
        or stripped.startswith("|")
    )


def _top_level_section_units(text: str) -> list[ExtractedUnit]:
    matches = list(re.finditer(r"(?m)^Section\s+(\d+)\s*:\s*(.+)$", text))
    if not matches:
        return []

    units: list[ExtractedUnit] = []
    front_matter = text[: matches[0].start()].strip()
    if front_matter:
        units.append(
            ExtractedUnit(
                id="front-matter",
                text=front_matter,
                unit_type="front_matter",
                metadata={"title": "Front Matter", "level": 0},
            )
        )

    for match_index, match in enumerate(matches):
        section_number = int(match.group(1))
        section_start = match.start()
        next_start = matches[match_index + 1].start() if match_index + 1 < len(matches) else len(text)
        section_text = text[section_start:next_start].strip()
        section_lines = section_text.splitlines()
        title_parts = [match.group(2).strip()]

        for line in section_lines[1:4]:
            stripped = line.strip()
            if not stripped:
                continue
            if re.match(rf"^{section_number}\.\d+\b", stripped):
                break
            if re.match(r"^Section\s+\d+\b", stripped):
                break
            if re.match(r"^Page\s+\d+\b", stripped):
                break
            title_parts.append(stripped)

        title = " ".join(title_parts)
        title = re.sub(r"\s+", " ", title).strip()
        units.append(
            ExtractedUnit(
                id=f"section-{section_number}",
                text=section_text,
                unit_type="section",
                metadata={"title": title, "level": 1, "section": section_number},
            )
        )

    return units


def _pdf_section_units(content: bytes) -> tuple[str, list[ExtractedUnit]] | None:
    try:
        import fitz

        doc = fitz.open(stream=content, filetype="pdf")
    except Exception:
        return None

    blocks: list[_PdfTextBlock] = []
    page_texts: list[str] = []
    table_areas = _pdf_table_areas(content)
    for page_index, page in enumerate(doc, start=1):
        page_texts.append(page.get_text("text").strip())
        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:
                continue

            parts: list[str] = []
            sizes: list[float] = []
            for line in block.get("lines", []):
                line_parts = []
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if text:
                        line_parts.append(text)
                        sizes.append(float(span.get("size", 0)))
                if line_parts:
                    parts.append(" ".join(line_parts))

            text = re.sub(r"\s+", " ", " ".join(parts)).strip()
            if text:
                bbox = tuple(round(float(value), 2) for value in block.get("bbox", (0, 0, 0, 0)))
                blocks.append(
                    _PdfTextBlock(
                        text=text,
                        page=page_index,
                        bbox=bbox,
                        max_size=max(sizes or [0.0]),
                        is_table=_overlaps_any_table(bbox, table_areas.get(page_index, [])),
                    )
                )

    full_text = "\n\n".join(page_text for page_text in page_texts if page_text)
    repeated_texts = _repeated_pdf_texts(blocks)
    body_size = _dominant_body_size(blocks)
    headings = [
        (index, block)
        for index, block in enumerate(blocks)
        if _is_pdf_heading(block, body_size)
    ]
    headings = _filter_repeated_prominent_titles(headings, body_size)
    if not headings:
        return full_text, []

    units: list[ExtractedUnit] = []
    for heading_number, (heading_index, heading) in enumerate(headings, start=1):
        next_heading_index = headings[heading_number][0] if heading_number < len(headings) else len(blocks)
        section_blocks = [
            block
            for block in blocks[heading_index:next_heading_index]
            if not _is_repeated_page_footer(block) and not block.is_table and _normalized_heading(block.text) not in repeated_texts
        ]
        section_text = "\n\n".join(block.text for block in section_blocks).strip()
        units.append(
            ExtractedUnit(
                id=f"section-{heading_number}",
                text=section_text,
                unit_type="section",
                metadata={
                    "title": heading.text,
                    "level": 1,
                    "section": heading_number,
                    "page": heading.page,
                    "bbox": heading.bbox,
                    "order": _position_order(heading.page, heading.bbox),
                    "reason": "pdf_heading",
                },
            )
        )

    return full_text, units


def _dominant_body_size(blocks: list[_PdfTextBlock]) -> float:
    sizes = [round(block.max_size, 1) for block in blocks if block.max_size >= 8]
    if not sizes:
        return 12.0
    return Counter(sizes).most_common(1)[0][0]


def _repeated_pdf_texts(blocks: list[_PdfTextBlock]) -> set[str]:
    pages_by_text: dict[str, set[int]] = {}
    for block in blocks:
        text = _normalized_heading(block.text)
        if not text or len(text) > 80:
            continue
        pages_by_text.setdefault(text, set()).add(block.page)
    return {text for text, pages in pages_by_text.items() if len(pages) >= 3}


def _filter_repeated_prominent_titles(
    headings: list[tuple[int, _PdfTextBlock]],
    body_size: float,
) -> list[tuple[int, _PdfTextBlock]]:
    prominent_titles = [
        _normalized_heading(block.text)
        for _, block in headings
        if block.max_size >= body_size + 3
    ]
    repeated = {title for title, count in Counter(prominent_titles).items() if title and count > 1}
    if not repeated:
        return headings
    return [
        (index, block)
        for index, block in headings
        if not (block.max_size >= body_size + 3 and _normalized_heading(block.text) in repeated)
    ]


def _normalized_heading(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip().casefold()


def _is_pdf_heading(block: _PdfTextBlock, body_size: float) -> bool:
    text = block.text.strip()
    if _is_repeated_page_footer(block):
        return False
    if len(text) > 90 or len(text.split()) > 8:
        return False
    if re.search(r"[.!?]$", text):
        return False
    if text.upper() in {"INDEX", "TABLE OF CONTENTS"}:
        return True
    if block.max_size >= body_size + 3:
        return bool(re.search(r"[A-Za-z]", text))
    return _is_structural_pdf_heading(block)


def _is_repeated_page_footer(block: _PdfTextBlock) -> bool:
    text = block.text.strip()
    return bool(
        re.match(r"^Page\s+\d+$", text, re.IGNORECASE)
        or re.match(r"^.+\s+\d+$", text)
        or re.match(r"^Version\s+[\d.]+$", text, re.IGNORECASE)
        or text.lower() == "internal"
    )


def _is_structural_pdf_heading(block: _PdfTextBlock) -> bool:
    text = block.text.strip()
    if block.page <= 2:
        return False
    if block.bbox[0] > 90 or block.bbox[1] < 70 or block.bbox[1] > 720:
        return False
    if re.match(r"^\(?[A-Z]\)\s+[A-Z][A-Z\s-]+$", text):
        return True
    if re.match(r"^\d+[.)]?\s+[A-Z][A-Z\s/&-]+$", text):
        return True
    letters = re.findall(r"[A-Za-z]", text)
    if not letters:
        return False
    uppercase_ratio = sum(1 for letter in letters if letter.isupper()) / len(letters)
    return uppercase_ratio >= 0.85 and len(text.split()) <= 7


class SectionWiseExtractor(BaseExtractor):
    method_id = "section_wise"
    name = "Section-wise Extraction"
    description = "Detects headings and preserves section-level boundaries."

    def extract(self, document: UploadedDocument) -> ExtractionResult:
        if document.suffix == ".pdf":
            pdf_result = _pdf_section_units(document.content)
            if pdf_result:
                text, pdf_units = pdf_result
                if pdf_units:
                    table_units = _extract_table_units(document, text)
                    units = _attach_unit_context(_ordered_units([*pdf_units, *table_units]))
                    return self._result(
                        text,
                        units,
                        {"sections": len(pdf_units), "tables_detected": len(table_units)},
                    )
        if document.suffix == ".docx":
            docx_result = _docx_ordered_units(document.content)
            if docx_result:
                text, units = docx_result
                sections = sum(1 for unit in units if unit.unit_type == "section")
                tables = sum(1 for unit in units if unit.unit_type == "table")
                return self._result(text, _attach_unit_context(units), {"sections": sections, "tables_detected": tables})

        text = self._read_text(document)
        table_units = _extract_table_units(document, text)
        section_text = _text_without_tables(document, text)
        units = _section_units(section_text)
        return self._result(
            "\n\n".join([section_text, *(unit.text for unit in table_units)]).strip(),
            _attach_unit_context(_ordered_units([*units, *table_units])),
            {"sections": len(units), "tables_detected": len(table_units)},
        )


class HierarchicalExtractor(BaseExtractor):
    method_id = "hierarchical"
    name = "Hierarchical Extraction"
    description = "Builds a lightweight hierarchy from detected headings."

    def extract(self, document: UploadedDocument) -> ExtractionResult:
        section_result = SectionWiseExtractor().extract(document)
        text = section_result.text
        sections = [unit for unit in section_result.units if unit.unit_type == "section"]
        roots: list[ExtractedUnit] = []
        stack: list[ExtractedUnit] = []
        for section in sections:
            level = int(section.metadata.get("level", 1))
            while stack and int(stack[-1].metadata.get("level", 1)) >= level:
                stack.pop()
            if stack:
                stack[-1].children.append(section)
                section.metadata["parent"] = stack[-1].id
            else:
                roots.append(section)
            stack.append(section)
        table_units = [unit for unit in section_result.units if unit.unit_type == "table"]
        return self._result(
            text,
            _attach_unit_context(_ordered_units([*roots, *table_units])),
            {"sections": len(sections), "hierarchy_roots": len(roots), "tables_detected": len(table_units)},
        )


class LayoutAwareExtractor(BaseExtractor):
    method_id = "layout_aware"
    name = "Layout-aware Extraction"
    description = "Preserves page, block, and bounding-box metadata when the parser exposes it."

    def extract(self, document: UploadedDocument) -> ExtractionResult:
        if document.suffix == ".docx":
            docx_result = _docx_block_units(document.content)
            if docx_result:
                text, units = docx_result
                return self._result(
                    text,
                    _attach_unit_context(units),
                    {
                        "layout_blocks": sum(1 for unit in units if unit.unit_type != "table"),
                        "tables_detected": sum(1 for unit in units if unit.unit_type == "table"),
                    },
                )

        blocks = self._read_pdf_blocks(document.content) if document.suffix == ".pdf" else []
        if not blocks:
            return SectionWiseExtractor().extract(document)
        units = [
            ExtractedUnit(
                id=f"layout-{idx}",
                text=str(block["text"]),
                unit_type="layout_block",
                metadata={k: v for k, v in block.items() if k != "text"},
            )
            for idx, block in enumerate(blocks, start=1)
        ]
        text = "\n\n".join(unit.text for unit in units)
        table_units = _extract_table_units(document, text)
        return self._result(
            text,
            _attach_unit_context(_ordered_units([*units, *table_units])),
            {"layout_blocks": len(units), "tables_detected": len(table_units)},
        )


class TableAwareExtractor(BaseExtractor):
    method_id = "table_aware"
    name = "Table-aware Extraction"
    description = "Extracts text plus simple table units from PDFs, DOCX, HTML, and markdown-like text."

    def extract(self, document: UploadedDocument) -> ExtractionResult:
        text = self._read_text(document)
        table_units = _extract_table_units(document, text)
        if document.suffix == ".pdf":
            pdf_result = _pdf_section_units(document.content)
            section_text, units = pdf_result if pdf_result else (text, _section_units(text))
            units = [*units, *table_units]
        elif document.suffix == ".docx":
            docx_result = _docx_ordered_units(document.content)
            if docx_result:
                section_text, units = docx_result
            else:
                section_text = _text_without_tables(document, text)
                units = _section_units(section_text)
                units = [*units, *table_units]
        else:
            section_text = _text_without_tables(document, text)
            units = _section_units(section_text)
            units = [*units, *table_units]
        units = _attach_unit_context(_ordered_units(units))
        return self._result(
            "\n\n".join(unit.text for unit in units).strip(),
            units,
            {"tables_detected": sum(1 for unit in units if unit.unit_type == "table")},
        )


def _ordered_units(units: list[ExtractedUnit]) -> list[ExtractedUnit]:
    return sorted(units, key=_unit_sort_key)


def _attach_unit_context(units: list[ExtractedUnit]) -> list[ExtractedUnit]:
    current_section_id = ""
    current_section_title = ""
    current_path: list[str] = []
    for unit in units:
        if unit.unit_type == "section":
            title = str(unit.metadata.get("title", "")).strip()
            level = max(1, int(unit.metadata.get("level", 1) or 1))
            current_path = current_path[: level - 1]
            if title:
                current_path.append(title)
            current_section_id = unit.id
            current_section_title = title
            unit.metadata["section_path"] = " > ".join(current_path)
        elif current_section_id:
            unit.metadata.setdefault("parent_section_id", current_section_id)
            unit.metadata.setdefault("parent_section_title", current_section_title)
            unit.metadata.setdefault("section_path", " > ".join(current_path))
    next_section_id = ""
    next_section_title = ""
    next_section_path = ""
    for unit in reversed(units):
        if unit.unit_type == "section":
            next_section_id = unit.id
            next_section_title = str(unit.metadata.get("title", "")).strip()
            next_section_path = str(unit.metadata.get("section_path", "")).strip()
        elif unit.unit_type == "table" and not unit.metadata.get("parent_section_id") and next_section_id:
            unit.metadata["parent_section_id"] = next_section_id
            unit.metadata["parent_section_title"] = next_section_title
            unit.metadata["section_path"] = next_section_path
    return units


def _unit_sort_key(unit: ExtractedUnit) -> tuple[float, int]:
    order = unit.metadata.get("order")
    if isinstance(order, (int, float)):
        return float(order), 0 if unit.unit_type != "table" else 1

    page = unit.metadata.get("page")
    bbox = unit.metadata.get("bbox")
    if isinstance(page, int) and _is_bbox(bbox):
        return _position_order(page, bbox), 0 if unit.unit_type != "table" else 1

    section = unit.metadata.get("section")
    if isinstance(section, int):
        return float(section), 0

    table_index = unit.metadata.get("table_index")
    if isinstance(table_index, int):
        return 1_000_000 + float(table_index), 1

    return 1_000_000_000, 1


def _position_order(page: int, bbox: tuple[float, float, float, float]) -> float:
    return (page * 10_000) + bbox[1]


def _is_bbox(value: object) -> bool:
    return (
        isinstance(value, tuple)
        and len(value) == 4
        and all(isinstance(item, (int, float)) for item in value)
    )


def _extract_table_units(document: UploadedDocument, text: str) -> list[ExtractedUnit]:
    if document.suffix == ".pdf":
        return _pdf_table_units(document.content)
    if document.suffix == ".docx":
        return _docx_table_units(document.content)
    if document.suffix in {".html", ".htm"}:
        return _html_table_units(document.content)
    if document.suffix in {".txt", ".md", ".markdown"}:
        return _delimited_table_units(text)
    return _delimited_table_units(text)


def _pdf_pages_without_tables(content: bytes) -> list[str]:
    table_areas = _pdf_table_areas(content)
    try:
        import fitz

        doc = fitz.open(stream=content, filetype="pdf")
        pages: list[str] = []
        for page_index, page in enumerate(doc, start=1):
            parts = []
            for block in page.get_text("blocks"):
                bbox = tuple(round(float(value), 2) for value in block[:4])
                text = normalize_whitespace(str(block[4]))
                if text and not _overlaps_any_table(bbox, table_areas.get(page_index, [])):
                    parts.append(text)
            pages.append("\n\n".join(parts).strip())
        return pages
    except Exception:
        try:
            import pdfplumber

            pages = []
            with pdfplumber.open(BytesIO(content)) as pdf:
                for page_index, page in enumerate(pdf.pages, start=1):
                    table_bboxes = table_areas.get(page_index, [])
                    filtered = page
                    for bbox in table_bboxes:
                        filtered = filtered.filter(lambda obj, table_bbox=bbox: not _pdfplumber_obj_overlaps_table(obj, table_bbox))
                    pages.append((filtered.extract_text() or "").strip())
            return pages
        except Exception:
            return []


def _pdfplumber_obj_overlaps_table(obj: dict[str, object], table_bbox: tuple[float, float, float, float]) -> bool:
    try:
        bbox = (
            float(obj["x0"]),
            float(obj["top"]),
            float(obj["x1"]),
            float(obj["bottom"]),
        )
    except Exception:
        return False
    return _overlap_ratio(bbox, table_bbox) > 0.15


def _docx_page_units(content: bytes, target_chars: int = 2600) -> tuple[str, list[ExtractedUnit]] | None:
    ordered = _docx_linear_units(content)
    if ordered is None:
        return None

    full_text = "\n\n".join(unit.text for unit in ordered).strip()
    pages: list[ExtractedUnit] = []
    page_parts: list[str] = []
    page_start_order = 1.0
    page_number = 1
    explicit_breaks = any(unit.metadata.get("page_break_after") for unit in ordered)

    def flush(force: bool = False) -> None:
        nonlocal page_parts, page_number, page_start_order
        page_text = "\n\n".join(part for part in page_parts if part.strip()).strip()
        if not page_text and not force:
            return
        if page_text:
            pages.append(
                ExtractedUnit(
                    id=f"page-{page_number}",
                    text=page_text,
                    unit_type="page",
                    metadata={
                        "page": page_number,
                        "order": page_start_order,
                        "page_source": "explicit_break" if explicit_breaks else "estimated_docx_page",
                    },
                )
            )
            page_number += 1
        page_parts = []
        page_start_order = float(page_number)

    for unit in ordered:
        unit_order = float(unit.metadata.get("order", page_number))
        if not page_parts:
            page_start_order = unit_order
        projected = len("\n\n".join([*page_parts, unit.text]))
        should_estimate_break = not explicit_breaks and page_parts and projected > target_chars
        if should_estimate_break:
            flush()
            page_start_order = unit_order
        page_parts.append(unit.text)
        if unit.metadata.get("page_break_after"):
            flush(force=True)
    flush()

    for unit in ordered:
        unit_page = _find_docx_unit_page(unit, pages)
        if unit_page:
            unit.metadata.setdefault("page", unit_page)
    return full_text, _ordered_units([*pages, *[unit for unit in ordered if unit.unit_type == "table"]])


def _find_docx_unit_page(unit: ExtractedUnit, pages: list[ExtractedUnit]) -> int | None:
    unit_text = unit.text.strip()
    if not unit_text:
        return None
    probe = unit_text[: min(120, len(unit_text))]
    for page in pages:
        if probe and probe in page.text:
            return int(page.metadata.get("page", 0) or 0)
    return None


def _docx_block_units(content: bytes) -> tuple[str, list[ExtractedUnit]] | None:
    ordered = _docx_linear_units(content)
    if ordered is None:
        return None
    return "\n\n".join(unit.text for unit in ordered).strip(), _ordered_units(ordered)


def _docx_linear_units(content: bytes) -> list[ExtractedUnit] | None:
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        from docx.oxml.ns import qn

        doc = Document(BytesIO(content))
        units: list[ExtractedUnit] = []
        table_count = 0
        paragraph_count = 0
        order = 0
        for child in doc.element.body.iterchildren():
            if child.tag == qn("w:p"):
                order += 1
                paragraph = Paragraph(child, doc)
                text = paragraph.text.strip()
                if not text and not _docx_has_page_break(child):
                    continue
                if text:
                    paragraph_count += 1
                    units.append(
                        ExtractedUnit(
                            id=f"paragraph-{paragraph_count}",
                            text=text,
                            unit_type="paragraph",
                            metadata={
                                "paragraph": paragraph_count,
                                "order": order,
                                "style": getattr(getattr(paragraph, "style", None), "name", "") or "",
                                "page_break_after": _docx_has_page_break(child),
                            },
                        )
                    )
                elif units and _docx_has_page_break(child):
                    units[-1].metadata["page_break_after"] = True
            elif child.tag == qn("w:tbl"):
                order += 1
                table = Table(child, doc)
                rows = _clean_table_rows([[cell.text for cell in row.cells] for row in table.rows])
                if not _is_valid_table(rows):
                    continue
                table_count += 1
                units.append(_make_table_unit(table_count, rows, "docx", {"order": order}))
        return units
    except Exception:
        return None


def _docx_has_page_break(element: object) -> bool:
    xml = getattr(element, "xml", "")
    return 'w:type="page"' in xml or "lastRenderedPageBreak" in xml


def _docx_ordered_units(content: bytes) -> tuple[str, list[ExtractedUnit]] | None:
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        from docx.oxml.ns import qn

        doc = Document(BytesIO(content))
        units: list[ExtractedUnit] = []
        full_parts: list[str] = []
        current_title = "Introduction"
        current_level = 1
        current_lines: list[str] = []
        section_count = 0
        table_count = 0
        order = 0
        body = doc.element.body

        def flush_section() -> None:
            nonlocal section_count, current_lines
            text = "\n\n".join(line for line in current_lines if line.strip()).strip()
            if not text:
                current_lines = []
                return
            section_count += 1
            units.append(
                ExtractedUnit(
                    id=f"section-{section_count}",
                    text=text,
                    unit_type="section",
                    metadata={
                        "title": current_title,
                        "level": current_level,
                        "section": section_count,
                        "order": order - 0.1,
                        "reason": "docx_heading",
                    },
                )
            )
            current_lines = []

        children = list(body.iterchildren())
        for child_index, child in enumerate(children):
            if child.tag == qn("w:p"):
                order += 1
                paragraph = Paragraph(child, doc)
                text = paragraph.text.strip()
                if not text:
                    continue
                full_parts.append(text)
                next_text = _next_docx_paragraph_text(children, child_index, doc)
                heading = _docx_heading(paragraph, next_text)
                if heading:
                    flush_section()
                    current_title, current_level = heading
                    current_lines = [text]
                else:
                    current_lines.append(text)
            elif child.tag == qn("w:tbl"):
                order += 1
                table = Table(child, doc)
                rows = _clean_table_rows([[cell.text for cell in row.cells] for row in table.rows])
                if not _is_valid_table(rows):
                    continue
                flush_section()
                table_count += 1
                table_unit = _make_table_unit(table_count, rows, "docx", {"order": order})
                units.append(table_unit)
                full_parts.append(table_unit.text)

        flush_section()
        return "\n\n".join(full_parts).strip(), _ordered_units(units)
    except Exception:
        return None


def _next_docx_paragraph_text(children: list[object], start_index: int, doc: object) -> str:
    try:
        from docx.text.paragraph import Paragraph
        from docx.oxml.ns import qn

        for child in children[start_index + 1 :]:
            if child.tag != qn("w:p"):
                continue
            text = Paragraph(child, doc).text.strip()
            if text:
                return text
    except Exception:
        return ""
    return ""


def _docx_heading(paragraph: object, next_text: str) -> tuple[str, int] | None:
    text = paragraph.text.strip()
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    match = re.search(r"heading\s*(\d+)", style_name, re.IGNORECASE)
    if match:
        return text, int(match.group(1))

    numbered = re.match(r"^(\d+(?:\.\d+)*)\s+(.{3,})$", text)
    if numbered:
        return numbered.group(2).strip(), numbered.group(1).count(".") + 1

    if _looks_like_text_heading(text, next_text):
        return text, 1
    return None


def _text_without_tables(document: UploadedDocument, text: str) -> str:
    if document.suffix == ".docx":
        try:
            from docx import Document

            doc = Document(BytesIO(document.content))
            return "\n\n".join(paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip())
        except Exception:
            return text
    if document.suffix in {".html", ".htm"}:
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(document.content.decode("utf-8", errors="ignore"), "html.parser")
            for table in soup.find_all("table"):
                table.decompose()
            return html_to_text(str(soup))
        except Exception:
            return text
    if document.suffix in {".txt", ".md", ".markdown"}:
        return _strip_delimited_table_blocks(text)
    return text


def _strip_delimited_table_blocks(text: str) -> str:
    kept_lines: list[str] = []
    current: list[str] = []
    current_kind: str | None = None

    def flush() -> None:
        nonlocal current, current_kind
        if not current or current_kind is None:
            current = []
            current_kind = None
            return
        rows = _parse_delimited_rows(current, current_kind)
        if not _is_valid_table(rows):
            kept_lines.extend(current)
        current = []
        current_kind = None

    for line in text.splitlines():
        kind = _delimited_line_kind(line)
        if kind is None:
            flush()
            kept_lines.append(line)
            continue
        if current_kind is not None and kind != current_kind:
            flush()
        current_kind = kind
        current.append(line)
    flush()
    return "\n".join(kept_lines).strip()


def _pdf_table_areas(content: bytes) -> dict[int, list[tuple[float, float, float, float]]]:
    try:
        logging.getLogger("pdfminer").setLevel(logging.ERROR)
        import pdfplumber

        areas: dict[int, list[tuple[float, float, float, float]]] = {}
        with pdfplumber.open(BytesIO(content)) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                page_areas = []
                for table in page.find_tables() or []:
                    rows = _clean_table_rows(table.extract())
                    bbox = tuple(round(float(value), 2) for value in table.bbox)
                    if _is_valid_pdf_table(rows, bbox):
                        page_areas.append(bbox)
                if page_areas:
                    areas[page_number] = page_areas
        return areas
    except Exception:
        return {}


def _overlaps_any_table(
    bbox: tuple[float, float, float, float],
    table_bboxes: list[tuple[float, float, float, float]],
) -> bool:
    return any(_overlap_ratio(bbox, table_bbox) > 0.15 for table_bbox in table_bboxes)


def _overlap_ratio(a: tuple[float, float, float, float], b: tuple[float, float, float, float]) -> float:
    left = max(a[0], b[0])
    top = max(a[1], b[1])
    right = min(a[2], b[2])
    bottom = min(a[3], b[3])
    if right <= left or bottom <= top:
        return 0.0
    intersection = (right - left) * (bottom - top)
    area = max((a[2] - a[0]) * (a[3] - a[1]), 1.0)
    return intersection / area


def _pdf_table_units(content: bytes) -> list[ExtractedUnit]:
    try:
        logging.getLogger("pdfminer").setLevel(logging.ERROR)
        import pdfplumber

        units: list[ExtractedUnit] = []
        with pdfplumber.open(BytesIO(content)) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                page_table_count = 0
                for table in page.find_tables() or []:
                    rows = _repair_pdf_table_rows(_clean_table_rows(table.extract()))
                    bbox = tuple(round(float(value), 2) for value in table.bbox)
                    if not _is_valid_pdf_table(rows, bbox):
                        continue
                    units.append(
                        _make_table_unit(
                            len(units) + 1,
                            rows,
                            "pdf",
                            {
                                "page": page_number,
                                "bbox": bbox,
                                "order": _position_order(page_number, bbox),
                            },
                        )
                    )
                    page_table_count += 1
                if page_table_count:
                    continue
                for rows in _pdf_text_strategy_tables(page):
                    units.append(
                        _make_table_unit(
                            len(units) + 1,
                            rows,
                            "pdf",
                            {
                                "page": page_number,
                                "order": (page_number * 10_000) + 9_000 + page_table_count,
                            },
                        )
                    )
                    page_table_count += 1
        return units
    except Exception:
        return []


def _pdf_text_strategy_tables(page: object) -> list[list[list[str]]]:
    tables: list[list[list[str]]] = []
    settings_options = [
        {
            "vertical_strategy": "text",
            "horizontal_strategy": "text",
            "snap_tolerance": 3,
            "join_tolerance": 3,
            "intersection_tolerance": 5,
            "min_words_vertical": 2,
            "min_words_horizontal": 1,
        },
        {
            "vertical_strategy": "lines",
            "horizontal_strategy": "text",
            "snap_tolerance": 3,
            "join_tolerance": 3,
            "intersection_tolerance": 5,
        },
    ]
    seen: set[str] = set()
    for settings in settings_options:
        try:
            extracted = page.extract_tables(table_settings=settings) or []
        except Exception:
            continue
        for table in extracted:
            rows = _repair_pdf_table_rows(_clean_table_rows(table))
            if not _is_valid_text_strategy_pdf_table(rows):
                continue
            key = "\n".join("|".join(row) for row in rows)
            if key in seen:
                continue
            seen.add(key)
            tables.append(rows)
    return tables


def _docx_table_units(content: bytes) -> list[ExtractedUnit]:
    try:
        from docx import Document

        doc = Document(BytesIO(content))
        units: list[ExtractedUnit] = []
        for table_index, table in enumerate(doc.tables, start=1):
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            rows = _clean_table_rows(rows)
            if _is_valid_table(rows):
                units.append(_make_table_unit(len(units) + 1, rows, "docx", {"order": 1_000_000 + table_index}))
        return units
    except Exception:
        return []


def _html_table_units(content: bytes) -> list[ExtractedUnit]:
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(content.decode("utf-8", errors="ignore"), "html.parser")
        units: list[ExtractedUnit] = []
        for table_index, table in enumerate(soup.find_all("table"), start=1):
            rows = []
            for tr in table.find_all("tr"):
                cells = [cell.get_text(" ", strip=True) for cell in tr.find_all(["th", "td"])]
                if cells:
                    rows.append(cells)
            rows = _clean_table_rows(rows)
            if _is_valid_table(rows):
                units.append(_make_table_unit(len(units) + 1, rows, "html", {"order": 1_000_000 + table_index}))
        return units
    except Exception:
        return []


def _delimited_table_units(text: str) -> list[ExtractedUnit]:
    units: list[ExtractedUnit] = []
    current: list[str] = []
    current_kind: str | None = None

    def flush() -> None:
        nonlocal current, current_kind
        if not current or current_kind is None:
            current = []
            current_kind = None
            return
        rows = _parse_delimited_rows(current, current_kind)
        if _is_valid_table(rows):
            units.append(_make_table_unit(len(units) + 1, rows, current_kind, {"order": len(units) + 1_000_000}))
        current = []
        current_kind = None

    for line in text.splitlines():
        kind = _delimited_line_kind(line)
        if kind is None:
            flush()
            continue
        if current_kind is not None and kind != current_kind:
            flush()
        current_kind = kind
        current.append(line)
    flush()
    return units


def _delimited_line_kind(line: str) -> str | None:
    stripped = line.strip()
    if not stripped:
        return None
    if "|" in stripped and len([cell for cell in stripped.strip("|").split("|") if cell.strip()]) >= 2:
        return "pipe"
    if "\t" in stripped and len([cell for cell in stripped.split("\t") if cell.strip()]) >= 2:
        return "tab"
    if (
        not re.match(r"^[-*+]\s+", stripped)
        and ", " not in stripped
        and stripped.count(",") >= 1
        and len([cell for cell in stripped.split(",") if cell.strip()]) >= 2
    ):
        return "csv"
    return None


def _parse_delimited_rows(lines: list[str], kind: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        stripped = line.strip()
        if kind == "pipe":
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            if cells and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells if cell.strip()):
                continue
        elif kind == "tab":
            cells = [cell.strip() for cell in stripped.split("\t")]
        else:
            cells = [cell.strip() for cell in stripped.split(",")]
        rows.append(cells)
    return _clean_table_rows(rows)


def _clean_table_rows(rows: list[list[object | None]]) -> list[list[str]]:
    cleaned_rows = []
    max_columns = max((len(row) for row in rows), default=0)
    for row in rows:
        cleaned = [_clean_table_cell(cell) for cell in row]
        cleaned.extend([""] * (max_columns - len(cleaned)))
        if any(cell for cell in cleaned):
            cleaned_rows.append(cleaned)
    return cleaned_rows


def _clean_table_cell(cell: object | None) -> str:
    text = "" if cell is None else str(cell)
    text = text.replace("\x00", ":")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _repair_pdf_table_rows(rows: list[list[str]]) -> list[list[str]]:
    if not rows:
        return rows
    column_count = max((len(row) for row in rows), default=0)
    if column_count != 3:
        return rows

    numeric_first_cells = sum(1 for row in rows if row and re.match(r"^\s*\d+\b", row[0]))
    if numeric_first_cells < max(3, len(rows) // 3):
        return rows

    repaired: list[list[str]] = []
    for row in rows:
        padded = row + [""] * (column_count - len(row))
        first = padded[0].strip()
        rest = [cell.strip() for cell in padded[1:] if cell.strip()]

        header_text = " ".join(cell for cell in padded if cell.strip())
        if re.search(r"\bsr\.?\s*no\.?\b", header_text, re.IGNORECASE):
            repaired.append(["SR. NO.", re.sub(r"\bsr\.?\s*no\.?\b", "", header_text, flags=re.IGNORECASE).strip() or "CONTENTS"])
            continue

        numbered = re.match(r"^(\d+)\s*(.*)$", first)
        if numbered:
            number = numbered.group(1)
            first_fragment = numbered.group(2).strip()
            content = _join_pdf_table_fragments([first_fragment, *rest])
            repaired.append([number, content])
        elif first or rest:
            repaired.append(["", _join_pdf_table_fragments([first, *rest])])

    return _clean_table_rows(repaired)


def _join_pdf_table_fragments(parts: list[str]) -> str:
    text = ""
    for part in [item.strip() for item in parts if item.strip()]:
        if not text:
            text = part
            continue

        first_token, _, remainder = part.partition(" ")
        previous_match = re.search(r"([A-Za-z]{1,8})$", text)
        previous_token = previous_match.group(1) if previous_match else ""

        if previous_token.isupper() and first_token.isupper() and (
            len(first_token) == 1 or len(previous_token) <= 3
        ):
            text = text[: -len(previous_token)] + previous_token + first_token
            if remainder:
                text = f"{text} {remainder}"
        elif (
            previous_token
            and previous_token[:1].isupper()
            and first_token.islower()
            and 2 <= len(first_token) <= 5
            and len(previous_token) <= 6
        ):
            text = text[: -len(previous_token)] + previous_token + first_token
            if remainder:
                text = f"{text} {remainder}"
        else:
            text = f"{text} {part}"

    return text


def _is_valid_table(rows: list[list[str]]) -> bool:
    if len(rows) < 2:
        return False
    column_count = max((len(row) for row in rows), default=0)
    if column_count < 2:
        return False
    filled_cells = sum(1 for row in rows for cell in row if cell)
    return filled_cells >= max(4, column_count + 1)


def _is_valid_text_strategy_pdf_table(rows: list[list[str]]) -> bool:
    if not _is_valid_pdf_table(rows):
        return False
    column_count = max((len(row) for row in rows), default=0)
    if column_count > 5:
        return False

    filled_rows = [[cell for cell in row if cell.strip()] for row in rows]
    if not filled_rows:
        return False
    dense_rows = sum(1 for row in filled_rows if len(row) >= max(2, column_count - 1))
    if dense_rows / len(filled_rows) < 0.55:
        return False

    filled_cells = [cell.strip() for row in rows for cell in row if cell.strip()]
    short_fragments = [
        cell
        for cell in filled_cells
        if re.fullmatch(r"[A-Za-z]{1,3}", cell) or re.fullmatch(r"[a-z]{4,8}", cell)
    ]
    if len(short_fragments) / len(filled_cells) > 0.35:
        return False

    header = [cell.strip() for cell in rows[0] if cell.strip()]
    return len(header) >= 2 and any(len(cell.split()) > 1 or len(cell) >= 6 for cell in header)


def _is_valid_pdf_table(rows: list[list[str]], bbox: tuple[float, float, float, float] | None = None) -> bool:
    if not _is_valid_table(rows):
        return False
    if bbox and _table_bbox_is_too_large(bbox):
        return False

    column_count = max((len(row) for row in rows), default=0)
    if column_count >= 6 and len(rows) >= 10:
        return False

    filled_cells = [cell.strip() for row in rows for cell in row if cell.strip()]
    if not filled_cells:
        return False

    short_fragments = [
        cell
        for cell in filled_cells
        if re.fullmatch(r"[A-Za-z]{1,3}", cell) or re.fullmatch(r"[a-z]{4,8}", cell)
    ]
    if len(short_fragments) / len(filled_cells) > 0.25:
        return False

    if _looks_like_fragmented_prose_table(rows):
        return False

    first_row_text = " ".join(cell for cell in rows[0] if cell.strip())
    if first_row_text and first_row_text[:1].islower() and len(first_row_text.split()) > 5:
        return False

    prose_like_cells = [cell for cell in filled_cells if len(cell.split()) > 16]
    return len(prose_like_cells) / len(filled_cells) <= 0.2


def _looks_like_fragmented_prose_table(rows: list[list[str]]) -> bool:
    column_count = max((len(row) for row in rows), default=0)
    if column_count < 4 or len(rows) < 8:
        return False

    filled_cells = [cell.strip() for row in rows for cell in row if cell.strip()]
    if not filled_cells:
        return False

    fragment_cells = 0
    lower_start_cells = 0
    sentence_end_cells = 0
    for cell in filled_cells:
        words = cell.split()
        if len(words) <= 3 and re.search(r"[A-Za-z]", cell):
            fragment_cells += 1
        if cell[:1].islower():
            lower_start_cells += 1
        if re.search(r"[.!?:;]$", cell):
            sentence_end_cells += 1

    fragment_ratio = fragment_cells / len(filled_cells)
    lower_start_ratio = lower_start_cells / len(filled_cells)
    sentence_end_ratio = sentence_end_cells / len(filled_cells)
    header = " ".join(cell for cell in rows[0] if cell.strip())
    generic_headers = sum(1 for cell in rows[0] if re.fullmatch(r"Column\s+\d+", cell.strip(), re.IGNORECASE))

    if generic_headers >= 2 and column_count >= 4 and len(rows) >= 6:
        return True

    return (
        fragment_ratio > 0.7
        and (lower_start_ratio > 0.25 or sentence_end_ratio > 0.25)
        and not re.search(r"\b(sr\.?\s*no|metric|target|shift|timing|break|contents)\b", header, re.IGNORECASE)
    )


def _table_bbox_is_too_large(bbox: tuple[float, float, float, float]) -> bool:
    _, top, _, bottom = bbox
    return bottom - top > 520


def _make_table_unit(index: int, rows: list[list[str]], source_format: str, metadata: dict[str, object]) -> ExtractedUnit:
    return ExtractedUnit(
        id=f"table-{index}",
        text=_table_to_markdown(rows),
        unit_type="table",
        metadata={
            "table_index": index,
            "format": source_format,
            "rows": rows,
            "row_count": len(rows),
            "column_count": max((len(row) for row in rows), default=0),
            **metadata,
        },
    )


def _table_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    column_count = max(len(row) for row in rows)
    padded = [row + [""] * (column_count - len(row)) for row in rows]
    header = padded[0]
    body = padded[1:]
    lines = [
        "| " + " | ".join(_escape_markdown_cell(cell) for cell in header) + " |",
        "| " + " | ".join("---" for _ in range(column_count)) + " |",
    ]
    for row in body:
        lines.append("| " + " | ".join(_escape_markdown_cell(cell) for cell in row) + " |")
    return "\n".join(lines)


def _escape_markdown_cell(cell: str) -> str:
    return cell.replace("|", "\\|")
